#!/usr/bin/env python3
"""Assemble le manuel complet Kajy 10e — version finale."""
import copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = "/home/claude/manuel-kajy-10e/output/"
SVG = "/home/claude/manuel-kajy-10e/svg/"
SCENES = "/home/claude/manuel-kajy-10e/scenes/"

# fichier -> image de mise en situation a inserer juste avant le bloc
IMAGE_MAP = {
    "Geo01_Ligne_Droite.docx": SCENES + "geo01-ligne.jpg",
    "Fiche_D_Addition_Soustraction.docx": SCENES + "aricu-piece.jpg",
    "AriU9_AddSub_200_500.docx": SCENES + "ariu9-marche.jpg",
    "Fiche_F_Multiplication.docx": SCENES + "fichef-oeufs.jpg",
    "Dup03_Multiplication_Partage.docx": SCENES + "dup03-crayons.jpg",
    "Dup08_Division_3Chiffres.docx": SCENES + "dup08-famille.jpg",
    "AriU8_Calcul_Prix.docx": SCENES + "ariu8-marche.jpg",
    "Geo15_Rectangle.docx": SCENES + "geo1517-facade.jpg",
    "Geo20_Perimetre_Rectangle.docx": SCENES + "geo2024-cloture.jpg",
    "Geo46_Triangles_Consolidation.docx": SCENES + "geo46-toit.jpg",
    "Mes02_Longueur_Metres.docx": SCENES + "mes02-metre.jpg",
    "Dup07_Poids.docx": SCENES + "mes34-balance.jpg",
    "Mes35_Lire_Heure.docx": SCENES + "mes35-horloge.jpg",
    "Dup04_Monnaie_100_200.docx": SCENES + "dup0406-comptoir.jpg",
    "Fiche_Corrigee_0-100_v2.docx": SCENES + "fichea-billes.jpg",
    "Fiche_B_0-1000.docx": SCENES + "ficheb-mangues.jpg",
    "Fiche_E_Jusqu10000.docx": SCENES + "fichee-riz.jpg",
}

ORDER = [
    ("part", "PARTIE 1 — NUMÉRATION"),
    ("file", "Fiche_Corrigee_0-100_v2.docx"),
    ("file", "Fiche_B_0-1000.docx"),
    ("file", "Fiche_E_Jusqu10000.docx"),
    ("file", "Dup01_Ampolony_Ambiny.docx"),
    ("file", "Dup02_Ordinal_Voisins.docx"),
    ("file", "Fiche_C_Caracteristiques.docx"),
    ("file", "AriU1_Ranger_En_Suites.docx"),
    ("file", "AriU2_Pair_Impair.docx"),
    ("part", "PARTIE 2 — ADDITION ET SOUSTRACTION"),
    ("file", "AriU9_AddSub_200_500.docx"),
    ("file", "Fiche_D_Addition_Soustraction.docx"),
    ("file", "AriU3_Strategies_AddSub_50.docx"),
    ("file", "Fiche_G_Strategies.docx"),
    ("file", "Dup09_Strategies_AddSub.docx"),
    ("part", "PARTIE 3 — MULTIPLICATION ET DIVISION"),
    ("file", "Dup03_Multiplication_Partage.docx"),
    ("file", "AriU4_Tables_2_4_8.docx"),
    ("file", "AriU10_Tables_3_6_9.docx"),
    ("file", "AriU11_Table_7.docx"),
    ("file", "AriU12_Tables_3_6_9_7_Liens.docx"),
    ("file", "AriU6_Tables_5_10.docx"),
    ("file", "Fiche_F_Multiplication.docx"),
    ("file", "Dup05_Division_Tables_3_6_9_7.docx"),
    ("file", "Dup08_Division_3Chiffres.docx"),
    ("file", "AriU5_Multiplication_Partage_Division.docx"),
    ("file", "AriU7_Application_Mult_Div.docx"),
    ("part", "PARTIE 4 — GÉOMÉTRIE"),
    ("file", "Geo01_Ligne_Droite.docx"),
    ("file", "Geo06_Angles_Droit_Aigu_Obtus.docx"),
    ("file", "Geo09_Angles_Renfort_Droit_Aigu.docx"),
    ("file", "Geo10_Angles_Renfort_Droit_Obtus.docx"),
    ("file", "Geo14_Construire_Mesurer_Angle.docx"),
    ("file", "Geo15_Rectangle.docx"),
    ("file", "Geo17_Carre.docx"),
    ("file", "Geo20_Perimetre_Rectangle.docx"),
    ("file", "Geo24_Perimetre_Carre.docx"),
    ("file", "Dup10_Triangles.docx"),
    ("file", "Geo46_Triangles_Consolidation.docx"),
    ("part", "PARTIE 5 — MESURES"),
    ("file", "Mes02_Longueur_Metres.docx"),
    ("file", "Mes05_Longueur_DM_CM.docx"),
    ("file", "Mes08_Longueur_M_DM_CM.docx"),
    ("file", "Dup07_Poids.docx"),
    ("file", "Mes34_Masses_Estimation.docx"),
    ("file", "Mes35_Lire_Heure.docx"),
    ("file", "Mes36_Duree_Heures.docx"),
    ("file", "Mes42_Calendrier_Relations.docx"),
    ("file", "Dup11_Calendrier.docx"),
    ("file", "Mes44_Duree_Jours_Annees.docx"),
    ("part", "PARTIE 6 — MONNAIE"),
    ("file", "Dup04_Monnaie_100_200.docx"),
    ("file", "Dup06_Monnaie_Echange.docx"),
    ("file", "Mes41_Billets_2000_5000.docx"),
    ("file", "AriU13_Conversion_Monnaie_1000.docx"),
    ("file", "AriU8_Calcul_Prix.docx"),
]
assert len([x for x in ORDER if x[0] == "file"]) == 50

master = Document()
section = master.sections[0]
section.top_margin = Cm(1.6); section.bottom_margin = Cm(1.6)
section.left_margin = Cm(1.8); section.right_margin = Cm(1.8)

def set_run_font(run, size=14, bold=False, color=None, italic=False):
    run.font.name = "Times New Roman"; run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)

def add_heading(text, size=18, color=(0xC0, 0x00, 0x00)):
    p = master.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(text), size, True, color)

def add_body(text, size=14, bold=False, italic=False, color=None, align=None):
    p = master.add_paragraph()
    if align: p.alignment = align
    set_run_font(p.add_run(text), size, bold, color, italic)
    return p

def add_title_page():
    for _ in range(6): master.add_paragraph()
    add_heading("MATHÉMATIQUES", 30, (0xC0,0,0))
    add_heading("Kilasy faha-10 (Kajy)", 20, (0x1F,0x4E,0x79))
    add_body("Manuel J-Lab Éditions — conforme au programme officiel", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body("Version V1", 12, align=WD_ALIGN_PARAGRAPH.CENTER)
    master.add_page_break()

def add_avant_propos():
    add_heading("AVANT-PROPOS")
    add_body("Ce manuel de mathématiques pour la classe de 10e (Kajy) transforme le programme officiel malgache en séances directement utilisables par l'enseignant. Il suit la progression des nombres, des opérations, de la géométrie, des mesures et de la monnaie telle que définie par le programme officiel FRA.")
    add_body("Chaque sous-thème comporte une fiche de préparation complète pour l'enseignant, suivie d'une page de leçon illustrée pour les élèves, puis d'exercices notés sur 20 points avec leur corrigé. Certaines fiches couvrent plusieurs séances du programme portant sur le même objectif : l'enseignant les adapte alors à chaque semaine, comme indiqué dans le tableau des paliers de chaque fiche concernée.")
    add_body("Le vocabulaire malgache est limité à l'encadré VOAMBOLANA de chaque leçon, et provient uniquement des termes du programme officiel ou de termes confirmés avec l'enseignante référente du projet.")
    master.add_page_break()

def add_mode_emploi():
    add_heading("COMMENT UTILISER CE MANUEL")
    add_body("L'enseignant suit la fiche de préparation et utilise les supports indiqués dans la colonne « Support et Matériel ». La page LEÇON est présentée aux élèves avec ses schémas et exemples. Les exercices d'Application et d'Évaluation de la fiche sont des exercices de classe ; seuls les exercices placés après la leçon, dans la section EXERCICES, sont notés sur 20 points.")
    add_body("Pour les fiches réutilisables (couvrant plusieurs séances), un tableau des paliers indique la semaine et la borne numérique visée : l'enseignant choisit alors, dans les exercices d'Application et d'Évaluation, le bloc correspondant au palier de la semaine en cours.")
    add_body("Les outils visuels D/U (dizaine/unité), C/D/U (centaine/dizaine/unité) et M/C/D/U (millier/centaine/dizaine/unité) sont rassemblés en annexe pour référence rapide, de même que 5 sujets d'examen types et un rappel des figures géométriques étudiées.")
    master.add_page_break()

def add_sommaire():
    add_heading("SOMMAIRE")
    for tag, val in ORDER:
        if tag == "part": add_body(val, 14, True)
    add_body("ANNEXES — Tableaux de numération, sujets d'examen, figures géométriques", 14, True)
    master.add_page_break()

def add_annexe_numeration():
    add_heading("ANNEXE 1 — Tableaux de numération")
    add_body("Ces trois tableaux résument les outils de décomposition des nombres utilisés tout au long du manuel : D/U (nombres jusqu'à 100), C/D/U (jusqu'à 1000), M/C/D/U (jusqu'à 10 000).")
    for img, cap in [
        (SVG+"DU-FA-47.png", "D/U — exemple : 47 = 4 dizaines et 7 unités"),
        (SVG+"CDU-348.png", "C/D/U — exemple : 348 = 3 centaines, 4 dizaines et 8 unités"),
        (SVG+"MCDU-3456.png", "M/C/D/U — exemple : 3456 = 3 milliers, 4 centaines, 5 dizaines et 6 unités"),
    ]:
        master.add_picture(img, width=Cm(9))
        master.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp = add_body(cap, 11, align=WD_ALIGN_PARAGRAPH.CENTER); master.add_paragraph()
    master.add_page_break()

def add_annexe_geometrie():
    add_heading("ANNEXE 2 — Les figures géométriques étudiées en 10e")
    shapes = [
        ("Ligne droite", "Une ligne qui ne change jamais de direction.", None),
        ("Angle droit, aigu, obtus", "Angle droit = ouverture du coin d'une feuille (vérifié à l'équerre). Aigu = plus fermé. Obtus = plus ouvert.", SVG+"angles-comparaison.png"),
        ("Rectangle", "4 côtés, 4 angles droits, côtés opposés égaux deux à deux.", None),
        ("Carré", "4 côtés égaux, 4 angles droits (cas particulier du rectangle).", None),
        ("Triangle", "3 côtés, 3 sommets.", None),
    ]
    for nom, desc, img in shapes:
        add_body(nom, 15, True, color=(0x1F,0x4E,0x79))
        add_body(desc, 13)
        if img:
            master.add_picture(img, width=Cm(11)); master.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        master.add_paragraph()
    add_body("Formules de périmètre (seules formules vues au programme) :", 14, True)
    add_body("Périmètre du rectangle = (longueur + largeur) × 2", 13)
    add_body("Périmètre du carré = côté × 4", 13)
    master.add_page_break()

# ============ 5 SUJETS D'EXAMEN ============
EXAMS = [
    {"titre": "SUJET 1 — Addition (nombres 0 à 100)",
     "items": [
        ("1. Écrivez en lettres (4 pts)", ["34", "78"]),
        ("2. Écrivez en chiffres (4 pts)", ["quarante-sept", "quatre-vingt-douze"]),
        ("3. Effectuez les opérations suivantes (4 pts)", ["45 + 23", "38 + 19"]),
        ("4. Tracez (4 pts)", ["un rectangle de 6 cm sur 3 cm", "un carré de 4 cm de côté"]),
        ("5. Problème (4 pts)", [
            "Voa mesure deux bouts de ficelle de 45 cm et 23 cm. Elle les attache bout à bout. Quelle est la longueur totale ?",
            "Un tissu mesure 68 cm. On en coupe 19 cm. Quelle longueur reste-t-il ?"]),
     ],
     "corrige": ["34 → trente-quatre ; 78 → soixante-dix-huit", "47 ; 92", "68 ; 57", "(tracés à vérifier)", "68 cm ; 49 cm"]},
    {"titre": "SUJET 2 — Soustraction (nombres 0 à 500, avec monnaie)",
     "items": [
        ("1. Écrivez en lettres (4 pts)", ["235", "160"]),
        ("2. Écrivez en chiffres (4 pts)", ["trois cent quatre-vingts", "deux cent quarante-cinq"]),
        ("3. Effectuez les opérations suivantes (4 pts)", ["380 − 145", "420 − 260"]),
        ("4. Tracez (4 pts)", ["un rectangle de 8 cm sur 5 cm", "un carré de 6 cm de côté"]),
        ("5. Problème (4 pts)", [
            "Bao a un billet de 500 Ariary. Elle achète un cahier à 145 Ariary. Combien lui reste-t-il ?",
            "Rina a 380 Ariary. Il dépense 260 Ariary. Combien lui reste-t-il ?"]),
     ],
     "corrige": ["235 → deux cent trente-cinq ; 160 → cent soixante", "380 ; 245", "235 ; 160", "(tracés à vérifier)", "355 Ar ; 120 Ar"]},
    {"titre": "SUJET 3 — Multiplication (tables de 3, 6, 9, 7, avec périmètre)",
     "items": [
        ("1. Écrivez en lettres (4 pts)", ["42", "72"]),
        ("2. Écrivez en chiffres (4 pts)", ["soixante-trois", "quarante-neuf"]),
        ("3. Effectuez les opérations suivantes (4 pts)", ["6 × 7", "9 × 8"]),
        ("4. Tracez (4 pts)", ["un carré de 5 cm de côté", "un rectangle de 7 cm sur 4 cm"]),
        ("5. Problème (4 pts)", [
            "Un jardin carré mesure 8 m de côté. Calcule son périmètre.",
            "Une classe range 6 rangées de 7 chaises. Combien de chaises en tout ?"]),
     ],
     "corrige": ["42 → quarante-deux ; 72 → soixante-douze", "63 ; 49", "42 ; 72", "(tracés à vérifier)", "32 m ; 42 chaises"]},
    {"titre": "SUJET 4 — Division (avec partage)",
     "items": [
        ("1. Écrivez en lettres (4 pts)", ["182", "107"]),
        ("2. Écrivez en chiffres (4 pts)", ["deux cent douze", "trois cent vingt-trois"]),
        ("3. Effectuez les opérations suivantes (4 pts)", ["728 ÷ 4", "856 ÷ 8"]),
        ("4. Tracez (4 pts)", ["un triangle", "un rectangle de 9 cm sur 4 cm"]),
        ("5. Problème (4 pts)", [
            "Une coopérative partage 728 kg de riz entre 4 familles à parts égales. Combien chaque famille reçoit-elle ?",
            "856 mangues sont réparties dans 8 caisses égales. Combien de mangues par caisse ?"]),
     ],
     "corrige": ["182 → cent quatre-vingt-deux ; 107 → cent sept", "212 ; 323", "182 ; 107", "(tracés à vérifier)", "182 kg ; 107 mangues"]},
    {"titre": "SUJET 5 — Mélange (scénario monétaire complet)",
     "items": [
        ("1. Écrivez en lettres (4 pts)", ["1400", "4500"]),
        ("2. Écrivez en chiffres (4 pts)", ["mille sept cent cinquante", "neuf cents"]),
        ("3. Effectuez les opérations suivantes (4 pts)", ["350 × 4", "4500 ÷ 5"]),
        ("4. Tracez (4 pts)", ["un carré de 7 cm de côté", "un rectangle de 10 cm sur 6 cm"]),
        ("5. Problème — deux opérations (4 pts)", [
            "Fara achète 4 cahiers à 350 Ariary chacun avec un billet de 2000 Ariary. Combien la vendeuse doit-elle lui rendre ?",
            "Une coopérative partage 4500 Ariary de bénéfice entre 5 membres à parts égales, puis ajoute une prime de 850 Ariary à chacun. Combien chaque membre reçoit-il en tout ?"]),
     ],
     "corrige": ["1400 → mille quatre cents ; 4500 → quatre mille cinq cents", "1750 ; 900", "1400 ; 900", "(tracés à vérifier)", "600 Ar ; 1750 Ar"]},
]

def add_annexe_examens():
    add_heading("ANNEXE 3 — 5 sujets d'examen types")
    add_body("Sujets progressifs, calés sur l'avancement du programme au fil de l'année. Chaque sujet est noté sur 20 points.", 12, italic=True)
    master.add_page_break()
    for ex in EXAMS:
        add_heading(ex["titre"], 16, (0x1F,0x4E,0x79))
        add_body("Noté sur 20 points", 12, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        for label, items in ex["items"]:
            add_body(label, 13, True)
            for i, it in enumerate(items, 1):
                add_body(f"{i}) {it}", 13)
        master.add_page_break()
    add_heading("CORRIGÉS DES 5 SUJETS", 18, (0xC0,0,0))
    for i, ex in enumerate(EXAMS, 1):
        add_body(f"Corrigé — {ex['titre']}", 14, True)
        for j, c in enumerate(ex["corrige"], 1):
            p = master.add_paragraph()
            set_run_font(p.add_run(f"{j}) "), 13)
            set_run_font(p.add_run(c), 13, True, (0xC2,0x18,0x5B))
    master.add_page_break()

def copy_body_from(path):
    src = Document(path)
    body = master.element.body
    sectPr = body.find(qn('w:sectPr'))
    for child in list(src.element.body):
        if child.tag == qn('w:sectPr'): continue
        if child.findall('.//' + qn('w:drawing')): continue
        new_child = copy.deepcopy(child)
        if sectPr is not None: sectPr.addprevious(new_child)
        else: body.append(new_child)

# ============ ASSEMBLAGE ============
add_title_page()
add_avant_propos()
add_mode_emploi()
add_sommaire()
add_annexe_numeration()
add_annexe_geometrie()
add_annexe_examens()

file_count = 0
first_file_in_part = True
for tag, val in ORDER:
    if tag == "part":
        master.add_page_break()
        add_heading(val, size=20, color=(0x1F, 0x4E, 0x79))
        first_file_in_part = True
    else:
        if not first_file_in_part:
            master.add_page_break()
        first_file_in_part = False
        if val in IMAGE_MAP:
            master.add_picture(IMAGE_MAP[val], width=Cm(11))
            master.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        copy_body_from(BASE + val)
        file_count += 1
        print(f"[{file_count}/50] {val} fusionné")

master.save("/home/claude/manuel-kajy-10e/output/Manuel_Kajy_10e_FINAL.docx")
print("\nManuel final généré.")
