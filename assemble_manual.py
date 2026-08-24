#!/usr/bin/env python3
"""
Assemble the enhanced Manuel_Calcul_9e_JLearn.docx:
- Injects pedagogical and geometric illustrations in lesson blocks before EXERCICES.
- Adds Annexe 1: Boîte à outils & Méthodes visuelles (Flèches fractions, conversions, triangles magiques, 4 étapes).
- Adds Annexe 2: Le Coin des Petits Génies (10 Jeux mathématiques & Défis psychotechniques + Corrigés).
- Adds Annexe 3: Index & Mémento complet des formules et notions clés (5 catégories détaillées).
- Adds Annexe 4: 10 Sujets d'examen officiels blancs (Opérations /20 + Problème /40) avec leurs 10 corrigés détaillés en tableau à 3 colonnes (Solutions | Résultats | Opérations).
- Updates Table des Matières.
"""

import os
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

import exam_content
import games_content
import tools_content

COLOR_TITLE = RGBColor(0xC0, 0x00, 0x00)       # Red #C00000
COLOR_SUBTITLE = RGBColor(0x1E, 0x7B, 0x34)    # Green #1E7B34
COLOR_KEYWORD = RGBColor(0x1F, 0x4E, 0x79)     # Deep Blue #1F4E79
COLOR_CORRIGE = RGBColor(0xC2, 0x18, 0x5B)     # Bordeaux/Rose #C2185B
COLOR_TEXT = RGBColor(0x22, 0x22, 0x22)        # Dark grey/black

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = COLOR_TITLE
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = COLOR_SUBTITLE
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = COLOR_KEYWORD
    return p

def add_paragraph_styled(doc, text, bold_prefix="", color=COLOR_TEXT, italic=False, size=Pt(10), space_after=Pt(3)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Times New Roman"
        r_pre.font.size = size
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_KEYWORD
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = size
    r.font.italic = italic
    r.font.color.rgb = color
    return p

def insert_image_centered(doc, image_path, width_cm=14.0):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run()
        run.add_picture(image_path, width=Cm(width_cm))

def build():
    print("Loading base document...")
    doc = Document("Manuel_Calcul_9e_JLearn_ORIGINAL_BACKUP.docx")
    
    print(f"Original document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables.")

    # 1. Update Table des matières at the top
    print("Updating Table des matières at the top of the manual...")
    for i, p in enumerate(doc.paragraphs[:120]):
        if "Annexe — Évaluations format examen" in p.text:
            p.text = "Annexe 1 — Boîte à outils & Méthodes visuelles"
        elif "Annexe — Formules et théorèmes" in p.text:
            p.text = "Annexe 2 — Le Coin des Petits Génies (Jeux & Défis psychotechniques)\n" \
                     "Annexe 3 — Index & Mémento complet des formules et notions clés\n" \
                     "Annexe 4 — 10 Sujets d'examen officiels blancs et corrigés détaillés"

    # 2. Inject geometric figures directly into the core lessons (before EXERCICES)
    lesson_illustrations = [
        (223, "assets_new/schema_rectangle_cote.png", 13.0),
        (263, "assets_new/tableau_conversion_fleches.png", 14.5),
        (440, "assets_new/schema_carre_quadrillage.png", 13.0),
        (603, "assets_new/schema_triangle_rectangle.png", 13.0),
        (788, "assets_new/methode_fleches_fractions.png", 14.5),
        (840, "assets_new/schema_cercle_elements.png", 13.0),
        (1000, "assets_new/triangles_magiques_economie.png", 14.0),
        (1033, "assets_new/schema_cube_3d_patron.png", 14.0),
        (1153, "assets_new/triangles_magiques_economie.png", 14.0),
        (1518, "assets_new/schema_parallelepipede_3d_patron.png", 14.0),
        (1544, "assets_new/schema_cadran_horloge.png", 13.5),
        (1667, "assets_new/schema_triangle_rectangle.png", 13.0),
    ]

    print("Enriching core lessons with geometric & methodological figures...")
    for p_idx, img_path, w_cm in lesson_illustrations:
        if not os.path.exists(img_path):
            continue
        # Find EXERCICES paragraph after p_idx
        for k in range(p_idx, min(p_idx + 30, len(doc.paragraphs))):
            if "EXERCICES" in doc.paragraphs[k].text.upper():
                target_p = doc.paragraphs[k]
                new_p = doc.add_paragraph()
                new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                new_p.paragraph_format.space_before = Pt(4)
                new_p.paragraph_format.space_after = Pt(6)
                r = new_p.add_run()
                r.add_picture(img_path, width=Cm(w_cm))
                target_p._element.addprevious(new_p._element)
                break

    # 3. Locate and clean up the old trailing placeholder annexes (at the very end, > 1700)
    cut_index = None
    for i, p in enumerate(doc.paragraphs):
        if i > 1700 and "Annexe — Évaluations format examen" in p.text:
            cut_index = i
            break
            
    if cut_index is not None:
        print(f"Trimming old placeholder annexes starting at paragraph {cut_index}...")
        for p in list(doc.paragraphs[cut_index:]):
            p._element.getparent().remove(p._element)
        if len(doc.tables) > 270:
            last_t = doc.tables[-1]
            if "Périmètre du rectangle" in last_t.rows[1].cells[0].text:
                print("Removing old table 273...")
                last_t._tbl.getparent().remove(last_t._tbl)

    # 4. Add Annexe 1 : Boîte à outils & Méthodes visuelles
    print("Generating Annexe 1: Boîte à Outils & Méthodes Visuelles...")
    doc.add_page_break()
    add_heading_1(doc, "ANNEXE 1 — BOÎTE À OUTILS & MÉTHODES VISUELLES")
    add_paragraph_styled(doc, "Cette section présente les repères visuels, les moyens mnémotechniques et les astuces méthodologiques indispensables pour maîtriser les calculs, les fractions, les conversions et la résolution de problèmes.", italic=True)

    # 1.1 Fractions papillon
    add_heading_2(doc, "1. La Méthode des Flèches Croisées pour les Fractions (Méthode Papillon)")
    add_paragraph_styled(doc, "Pour additionner ou soustraire deux fractions à dénominateurs différents, on utilise la technique du produit en croix et du dénominateur commun :", bold_prefix="Technique : ")
    add_paragraph_styled(doc, "• Flèche descendante (Haut gauche × Bas droit) : donne le premier numérateur.\n"
                              "• Flèche montante (Bas gauche × Haut droit) : donne le deuxième numérateur.\n"
                              "• Flèche horizontale basse (Bas gauche × Bas droit) : donne le dénominateur commun.")
    insert_image_centered(doc, "assets_new/methode_fleches_fractions.png", 14.5)

    # 1.2 Fraction d'une quantité
    add_heading_2(doc, "2. Comment Calculer la Fraction d'une Quantité ?")
    add_paragraph_styled(doc, "Règle d'or en deux flèches : on divise d'abord par le dénominateur (le bas) pour trouver la valeur d'une seule part, puis on multiplie par le numérateur (le haut) pour obtenir le total des parts voulues.", bold_prefix="Règle d'or : ")
    insert_image_centered(doc, "assets_new/methode_fraction_quantite.png", 14.5)

    # 1.3 Saut de puce conversions
    add_heading_2(doc, "3. Le Tableau de Conversion et la Règle des Sauts de Puce")
    add_paragraph_styled(doc, "Pour convertir sans jamais hésiter entre multiplier et diviser : vers la droite (->), chaque case ajoute un zéro (× 10) ; vers la gauche (<-), chaque case enlève un zéro ou décale la virgule (÷ 10).", bold_prefix="Sens des flèches : ")
    insert_image_centered(doc, "assets_new/tableau_conversion_fleches.png", 14.5)

    # 1.4 Triangles magiques
    add_heading_2(doc, "4. Les Triangles Magiques du Commerce et du Budget Familial")
    add_paragraph_styled(doc, "Le triangle de calcul permet de retrouver instantanément l'opération à effectuer (+ ou −) : la valeur du haut s'obtient par l'addition des deux cases du bas ; une valeur du bas s'obtient en soustrayant.", bold_prefix="Lecture visuelle : ")
    insert_image_centered(doc, "assets_new/triangles_magiques_economie.png", 14.5)

    # 1.5 4 étapes résolution problème
    add_heading_2(doc, "5. La Méthode Officielle de Résolution de Problèmes en 4 Étapes")
    add_paragraph_styled(doc, "Guide universel pour réussir l'épreuve de Problème à l'examen officiel :", bold_prefix="Méthodologie : ")
    insert_image_centered(doc, "assets_new/etape_resolution_probleme.png", 14.5)

    # 5. Add Annexe 2 : Le Coin des Petits Génies (Jeux & Défis Psychotechniques)
    print("Generating Annexe 2: Le Coin des Petits Génies (Jeux Mathématiques)...")
    doc.add_page_break()
    add_heading_1(doc, "ANNEXE 2 — LE COIN DES PETITS GÉNIES (JEUX & DÉFIS PSYCHOTECHNIQUES)")
    add_paragraph_styled(doc, "Ces 10 défis mathématiques et logiques sont spécialement conçus pour développer le calcul mental rapide, la déduction logique, l'agilité psychotechnique et le repérage géométrique.", italic=True)
    insert_image_centered(doc, "assets_new/schema_jeux_pyramide_carremagique.png", 14.0)

    for g in games_content.GAMES:
        add_heading_2(doc, g["num"])
        add_paragraph_styled(doc, g["intro"], italic=True)
        for item in g["items"]:
            add_paragraph_styled(doc, item, space_after=Pt(4))

    # Page break for games solutions
    doc.add_page_break()
    add_heading_1(doc, "CORRIGÉS DES JEUX ET DÉFIS PSYCHOTECHNIQUES")
    for sol_group in games_content.GAMES_SOLUTIONS:
        add_heading_2(doc, sol_group["num"])
        for item in sol_group["items"]:
            add_paragraph_styled(doc, item, space_after=Pt(4))

    # 6. Add Annexe 3 : Index & Mémento complet des formules et notions clés
    print("Generating Annexe 3: Index & Mémento Complet des Formules...")
    doc.add_page_break()
    add_heading_1(doc, "ANNEXE 3 — INDEX & MÉMENTO COMPLET DES FORMULES ET NOTIONS CLÉS")
    add_paragraph_styled(doc, "Ce mémento récapitule toutes les formules, propriétés géométriques, unités de mesure et règles économiques du programme officiel de 9e, avec leurs renvois précis aux blocs thématiques.", italic=True)

    for cat in tools_content.INDEX_CATEGORIES:
        add_heading_2(doc, cat["category"])
        add_paragraph_styled(doc, cat["description"], italic=True)

        # Create Table for this category
        tbl = doc.add_table(rows=1, cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl, color="1F4E79", sz="4")

        headers = [("Notion / Titre", Cm(4.0)), ("Formule / Règle", Cm(5.0)), ("Exemple d'application", Cm(5.5)), ("Bloc réf.", Cm(2.0))]
        hdr_row = tbl.rows[0]
        for idx, (h_text, w) in enumerate(headers):
            c = hdr_row.cells[idx]
            c.width = w
            set_cell_shading(c, "DDEEFF")
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h_text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = COLOR_KEYWORD

        for row_data in cat["rows"]:
            r_elem = tbl.add_row()
            for idx, cell_text in enumerate(row_data):
                c = r_elem.cells[idx]
                c.width = headers[idx][1]
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                if idx == 0:
                    r = p.add_run(cell_text)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(9)
                    r.font.bold = True
                    r.font.color.rgb = COLOR_TITLE
                elif idx == 3:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(cell_text)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(8.5)
                    r.font.bold = True
                    r.font.color.rgb = COLOR_SUBTITLE
                else:
                    r = p.add_run(cell_text)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(9)
                    r.font.color.rgb = COLOR_TEXT

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 7. Add Annexe 4 : 10 Sujets d'examen officiels blancs + Corrigés
    print("Generating Annexe 4: 10 Sujets d'Examen Officiels Blancs & Corrigés...")
    doc.add_page_break()
    add_heading_1(doc, "ANNEXE 4 — 10 SUJETS D'EXAMEN OFFICIELS BLANCS ET CORRIGÉS DÉTAILLÉS")
    add_paragraph_styled(doc, "Cette annexe réunit 10 sujets d'examen complets conformes au format officiel des examens scolaires à Madagascar (Partie 1 : Opérations sur 20 points + Partie 2 : Problème sur 40 points, Total sur 60 points).\n"
                              "Chaque sujet est suivi sur la page suivante de son corrigé officiel intégral, utilisant pour le problème le tableau modèle à trois colonnes : Solutions | Résultats | Opérations.", italic=True)

    for ex in exam_content.EXAMS:
        # 1) SUBJECT PAGE
        doc.add_page_break()
        add_heading_1(doc, f"{ex['title'].upper()}")
        add_paragraph_styled(doc, f"Niveau : Classe de 9e (CE)  |  Durée totale : {ex['duration']}  |  Barème officiel : {ex['total_points']} points (Opérations sur 20 + Problème sur 40)", bold_prefix="Format officiel : ", color=COLOR_KEYWORD)
        add_paragraph_styled(doc, f"Programme évalué : {ex['period']}", italic=True)

        # PART 1: OPÉRATIONS (/20 pts)
        add_heading_2(doc, "PREMIÈRE ÉPREUVE : OPÉRATIONS (20 points)")
        add_paragraph_styled(doc, "Effectue les calculs demandés avec soin et précision sur ta feuille de copie :", italic=True)

        for op in ex["operations"]:
            add_paragraph_styled(doc, op["consigne"], bold_prefix=f"{op['num']} — ")
            for it in op["items"]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(f"•  {it}")
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)

        # PART 2: PROBLÈME (/40 pts)
        add_heading_2(doc, f"DEUXIÈME ÉPREUVE : {ex['problem']['title']}")
        add_paragraph_styled(doc, ex['problem']['enonce'])
        add_heading_3(doc, "Questions à résoudre :")
        for q in ex['problem']['questions']:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"•  {q}")
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
            r.font.bold = True

        # 2) SOLUTION PAGE (PAGE BREAK)
        doc.add_page_break()
        add_heading_1(doc, f"CORRIGÉ OFFICIEL DÉTAILLÉ — {ex['title'].upper()}")
        add_paragraph_styled(doc, f"Barème total : {ex['total_points']} points  |  Seuil de réussite recommandé : 30 / 60 points", bold_prefix="Grille de notation : ", color=COLOR_CORRIGE)

        # Corrigé Opérations
        add_heading_2(doc, "1. Corrigé de l'Épreuve d'Opérations (20 points)")
        for c_op in ex["corriges_operations"]:
            add_paragraph_styled(doc, c_op, space_after=Pt(3))

        # Corrigé Problème en tableau à 3 colonnes
        add_heading_2(doc, "2. Corrigé de l'Épreuve de Problème (40 points) — Tableau officiel à 3 colonnes")
        add_paragraph_styled(doc, "Présentation canonique exigée aux examens officiels à Madagascar :", italic=True)

        # 3-column table
        table_sol = doc.add_table(rows=1, cols=3)
        table_sol.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table_sol, color="1F4E79", sz="6")

        hdr_cells = table_sol.rows[0].cells
        hdr_cells[0].width = Cm(7.5)
        hdr_cells[1].width = Cm(4.0)
        hdr_cells[2].width = Cm(4.5)

        headers_prob = [
            ("Solutions (Explications & Formules)", WD_ALIGN_PARAGRAPH.LEFT),
            ("Résultats (avec unité)", WD_ALIGN_PARAGRAPH.CENTER),
            ("Opérations (Posées)", WD_ALIGN_PARAGRAPH.RIGHT)
        ]
        for idx, (h_title, align) in enumerate(headers_prob):
            c = hdr_cells[idx]
            set_cell_shading(c, "DDEEFF")
            p = c.paragraphs[0]
            p.alignment = align
            r = p.add_run(h_title)
            r.font.name = "Times New Roman"
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = COLOR_KEYWORD

        for row_data in ex["problem_solution_rows"]:
            row_cells = table_sol.add_row().cells
            row_cells[0].width = Cm(7.5)
            row_cells[1].width = Cm(4.0)
            row_cells[2].width = Cm(4.5)

            # Solutions cell
            p0 = row_cells[0].paragraphs[0]
            p0.paragraph_format.space_before = Pt(3)
            p0.paragraph_format.space_after = Pt(3)
            r0 = p0.add_run(row_data["solution"] + f"  ({row_data['bareme']})")
            r0.font.name = "Times New Roman"
            r0.font.size = Pt(9.5)
            r0.font.bold = True

            # Results cell
            p1 = row_cells[1].paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.paragraph_format.space_before = Pt(3)
            p1.paragraph_format.space_after = Pt(3)
            r1 = p1.add_run(row_data["result"])
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(10)
            r1.font.bold = True
            r1.font.color.rgb = COLOR_SUBTITLE

            # Operations cell
            p2 = row_cells[2].paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p2.paragraph_format.space_before = Pt(2)
            p2.paragraph_format.space_after = Pt(2)
            r2 = p2.add_run(row_data["operation"])
            r2.font.name = "Courier New"
            r2.font.size = Pt(8.5)
            r2.font.bold = True
            r2.font.color.rgb = COLOR_CORRIGE

    output_path = "Manuel_Calcul_9e_JLearn.docx"
    print(f"Saving final enhanced document to {output_path}...")
    doc.save(output_path)
    print("Build finished successfully!")

if __name__ == "__main__":
    build()
