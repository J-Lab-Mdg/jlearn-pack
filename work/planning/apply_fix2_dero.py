# -*- coding: utf-8 -*-
"""
Fix #2 (déroulement tables): expand <4-item exercise groups to 4 items.
Uses element references (not positional indices) to stay valid across in-row insertions.
"""
import sys, re, copy
sys.path.insert(0, '/home/user/jlearn-pack/work/planning')
import content
import docx
from docx.oxml.ns import qn

SRC = '/home/user/jlearn-pack/work/manuel_v2.docx'

d = docx.Document(SRC)
tables6 = [t for t in d.tables if len(t.columns) == 6]

def parse_groups(paragraphs):
    """paragraphs: list of (elem, text). Returns list of dict with item elems/texts."""
    groups = []
    cur = None
    for el, text in paragraphs:
        t = text.strip()
        if not t:
            continue
        if re.match(r'^\d+\.\s', t):
            if cur is None:
                continue
            cur['item_elems'].append(el)
            cur['item_texts'].append(t)
        else:
            cur = {'consigne_text': t, 'item_elems': [], 'item_texts': []}
            groups.append(cur)
    return [g for g in groups if g['item_elems']]

def flat_items(paragraphs):
    out = []
    for el, text in paragraphs:
        t = text.strip()
        if re.match(r'^\d+\.\s', t):
            out.append((el, t))
    return out

total_edits = 0
log = []

for si, t in enumerate(tables6, start=1):
    for r in t.rows:
        label = r.cells[0].text.strip()
        if not (label == '6. Application' or label.startswith('III.')):
            continue
        ens_cell = r.cells[1]
        app_cell = r.cells[2]

        ens_paras = [(p._p, p.text) for p in ens_cell.paragraphs]
        app_paras = [(p._p, p.text) for p in app_cell.paragraphs]

        ens_groups = parse_groups(ens_paras)
        app_flat = flat_items(app_paras)

        offsets = []
        run = 0
        for g in ens_groups:
            offsets.append(run)
            run += len(g['item_elems'])

        for gi, g in enumerate(ens_groups):
            n_items = len(g['item_elems'])
            if n_items >= 4:
                continue
            sig = tuple(g['item_texts'])
            if sig not in content.DATA:
                raise SystemExit(f"NO CONTENT for S{si} {label} group {gi}: {sig}")
            spec = content.DATA[sig]
            new_ens_lines = spec['new_ens']
            new_app_lines = spec['new_app']
            need = 4 - n_items
            assert len(new_ens_lines) == need, f"S{si} {label} exo{gi}: need {need} got {len(new_ens_lines)}"
            assert len(new_app_lines) == need

            # ---- ENS side ----
            last_ens_p = g['item_elems'][-1]
            anchor = last_ens_p
            for k, line in enumerate(new_ens_lines):
                new_num = n_items + k + 1
                new_p = copy.deepcopy(last_ens_p)
                runs = new_p.findall(qn('w:r'))
                for extra in runs[1:]:
                    new_p.remove(extra)
                t_el = runs[0].find(qn('w:t'))
                t_el.text = f"{new_num}. {line}"
                t_el.set(qn('xml:space'), 'preserve')
                anchor.addnext(new_p)
                anchor = new_p

            # ---- APP side ----
            start = offsets[gi]
            end = start + n_items
            slice_ = app_flat[start:end]
            last_app_p = slice_[-1][0]
            anchor2 = last_app_p
            for k, line in enumerate(new_app_lines):
                new_num = n_items + k + 1
                new_p = copy.deepcopy(last_app_p)
                runs = new_p.findall(qn('w:r'))
                if len(runs) >= 2:
                    for extra in runs[2:]:
                        new_p.remove(extra)
                    t0 = runs[0].find(qn('w:t'))
                    t0.text = f"{new_num}. "
                    t0.set(qn('xml:space'), 'preserve')
                    t1 = runs[1].find(qn('w:t'))
                    t1.text = line
                elif len(runs) == 1:
                    for extra in runs[1:]:
                        new_p.remove(extra)
                    t0 = runs[0].find(qn('w:t'))
                    t0.text = f"{new_num}. {line}"
                    t0.set(qn('xml:space'), 'preserve')
                else:
                    raise SystemExit(f"APP para has no runs: S{si} {label} exo{gi}")
                anchor2.addnext(new_p)
                anchor2 = new_p

            total_edits += 1
            log.append(f"S{si} {label[:20]} exo#{gi}: +{need} items")

print(f"Applied {total_edits} exercise-group expansions (déroulement tables).")
for l in log:
    print(" ", l)

d.save(SRC)
print("Saved.")
